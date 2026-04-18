#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AGÜ Mobile teknik raporu Word (.docx) üretir. python-docx gerekir."""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# Anket görselleri: önce proje içi assets/survey_charts, yoksa Cursor assets veya AGU_REPORT_ASSETS
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_SCRIPT_DIR)
_DEFAULT_LOCAL_ASSETS = os.path.join(_PROJECT_ROOT, "report_assets", "survey_charts")
_FALLBACK_CURSOR_ASSETS = "/Users/mustafabicer/.cursor/projects/Users-mustafabicer-projects-Ag-Mobile/assets"
ASSET_DIR = os.environ.get("AGU_REPORT_ASSETS") or (
    _DEFAULT_LOCAL_ASSETS
    if os.path.isdir(_DEFAULT_LOCAL_ASSETS)
    and any(f.endswith(".png") for f in os.listdir(_DEFAULT_LOCAL_ASSETS))
    else _FALLBACK_CURSOR_ASSETS
)

OUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "AGÜ Teknik Raporu.docx",
)

SURVEY_IMAGES: list[tuple[str, str]] = [
    (
        "Ekran_Resmi_2026-04-08_19.36.58-ab758d6d-d9ea-4f1f-b14d-bc4e4edb23b0.png",
        "Şekil 1 — Mobil uygulama ihtiyacına ilişkin görüş (31 yanıt).",
    ),
    (
        "Ekran_Resmi_2026-04-08_19.37.07-012f7442-faa7-4795-8fcb-53366deb3192.png",
        "Şekil 2 — Resmî platform ve yönetim modeline ilişkin görüşler (31 yanıt).",
    ),
    (
        "Ekran_Resmi_2026-04-08_19.36.06-3095cb6a-5e76-4e2b-99c9-79466b0dbb17.png",
        "Şekil 3 — Uygulama memnuniyeti (31 yanıt).",
    ),
    (
        "Ekran_Resmi_2026-04-08_19.36.37-1c8407a7-863d-4ca0-8054-8d309ed55a83.png",
        "Şekil 4 — En çok kullanılan modüller (çoklu seçim, 31 yanıt).",
    ),
    (
        "Ekran_Resmi_2026-04-08_19.36.26-9c3b254f-e308-4935-b4f1-99db1346a5b5.png",
        "Şekil 5 — Öğrenci faydasına inanç (31 yanıt).",
    ),
    (
        "Ekran_Resmi_2026-04-08_19.36.16-d67e2bfe-cc06-4090-9f56-534ab95e0a66.png",
        "Şekil 6 — Kullanım sıklığı (31 yanıt).",
    ),
    (
        "Ekran_Resmi_2026-04-08_19.36.47-7090aebc-2b62-47a0-a32d-dd9dcdb7b24b.png",
        "Şekil 7 — Güvenlik ve gizlilik algısı (31 yanıt).",
    ),
]


def set_cell_shading(cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def add_page_break(paragraph) -> None:
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def style_body(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.bold = bold


def add_picture_safe(doc: Document, filename: str, caption: str) -> None:
    path = os.path.join(ASSET_DIR, filename)
    if not os.path.isfile(path):
        doc.add_paragraph(
            f"[Görsel dosyası bulunamadı: {filename} — İçeriği rapora elle ekleyebilirsiniz.]",
            style="Intense Quote",
        )
        return
    try:
        doc.add_picture(path, width=Inches(5.8))
    except Exception as e:
        doc.add_paragraph(f"[Görsel eklenemedi ({e}): {filename}]")
        return
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    style_body(doc)

    # ---- Kapak ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AGÜ MOBİLE\n")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x1A, 0x4A, 0x6E)
    t.add_run("\nTeknik ve Bilgilendirme Raporu\n").bold = True
    t.runs[-1].font.size = Pt(16)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "Abdullah Gül Üniversitesi\n"
        "Öğrenci Geliştirici Ekibi\n\n"
        "Hazırlayan: [Ekip üyelerinin adları buraya eklenecek]\n\n"
        f"Tarih: {date(2026, 4, 8).strftime('%d.%m.%Y')}\n"
        "Sürüm: 1.2 (uygulama sürümü ile uyumlu referans)"
    )
    mr.font.size = Pt(11)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Bu belge, üniversite yönetimine sunulmak üzere hazırlanmış resmî nitelikte "
        "bir teknik özet ve şeffaflık dokümanıdır."
    )
    nr.font.size = Pt(9)
    nr.italic = True

    add_page_break(doc.add_paragraph())

    # ---- İçindekiler ----
    h(doc, "İçindekiler", 1)
    toc_p = doc.add_paragraph()
    add_toc_field(toc_p)
    doc.add_paragraph()
    hint = doc.add_paragraph(
        "Not: Word’de içindekiler listesini güncellemek için tabloya sağ tıklayıp "
        "“Alanları Güncelle” seçeneğini kullanabilirsiniz."
    )
    hint.runs[0].italic = True
    hint.runs[0].font.size = Pt(9)
    add_page_break(doc.add_paragraph())

    # ---- Özet ----
    h(doc, "Yönetici Özeti", 1)
    body = """
Bu rapor, Abdullah Gül Üniversitesi öğrencileri tarafından geliştirilen AGÜ Mobile
uygulamasının teknik mimarisini, modüllerini, kişisel verilerin işlenmesi ve güvenlik
modelini; ayrıca kullanıcı geri bildirimlerini (anket) özetlemektedir. Uygulama,
öğrencilerin kampüs hayatını kolaylaştırmayı amaçlayan; ders programı, yemekhane,
etkinlikler, devamsızlık takibi ve üniversite sistemlerine hızlı erişim gibi işlevler
sunan çok platformlu (Android ve iOS) bir istemcidir.

Altyapıda Google Firebase (kimlik doğrulama, bulut veritabanı ve dosya depolama)
kullanılmaktadır. Kullanıcı parolaları geliştiriciler tarafından görülemez; kimlik
bilgileri endüstri standardı şekilde Google altyapısında yönetilir. Ders programının
SIS üzerinden alınması, kullanıcının tarayıcıda oturum açıp programı görüntülemesiyle
aynı veriyi cihazda düzenli biçimde saklamayı hedefleyen, ek bir “arka kapı” veya
veri sızıntısı içermeyen istemci tarafı bir süreçtir.

Geliştirici ekibin stratejik niyeti; uygulamanın olgunlaştırılması ve üniversite
yönetimine kontrol ile sorumluluğun tamamen okula devredileceği şekilde teslim
edilmesidir. Üniversitenin uygun görmesi halinde uygulamanın resmî bir kampüs
platformuna dönüştürülmesi hedeflenmektedir. Google Play’de yayımlanmış sürüm
hakkında geçmişte yaşanan iletişim eksikliği bu raporda açık ve ölçülü bir dille
anlatılmış; App Store sürümünün yalnızca üniversite onayıyla yayımlanması
kararlaştırılmıştır.
"""
    for line in body.strip().split("\n"):
        doc.add_paragraph(line.strip())

    h(doc, "Giriş ve Raporun Kapsamı", 1)
    doc.add_paragraph(
        "Bu doküman; uygulamanın işlevlerini modül modül özetler, gizlilik ve güvenlik "
        "konularında şeffaflık sağlar, SIS ile ders programı entegrasyonunun hukuki ve "
        "teknik çerçevesini ayrıntılandırır ve kullanıcı anket sonuçlarını sunar. "
        "Metin, teknik olmayan okuyucular için anlaşılır bir dil ile birlikte, inceleme "
        "yapacak bilgi işlem birimlerinin beklentilerini karşılayacak düzeyde teknik "
        "ayrıntı içerir."
    )

    h(doc, "Uygulama Hakkında Genel Bilgi", 1)
    doc.add_paragraph(
        "AGÜ Mobile, Flutter çatısı ile geliştirilmiştir. Bu sayede tek kod tabanından "
        "Android ve iOS sürümleri üretilebilmekte; arayüz tutarlılığı ve bakım maliyeti "
        "düşürülmektedir. Uygulama sürümü (pubspec referansı): 1.2.0+21."
    )
    doc.add_paragraph(
        "Kayıt sırasında kullanıcıya, uygulamanın üniversitenin resmî uygulaması "
        "olmadığı ve verilerin yalnızca uygulama işlevselliği için kullanılacağı "
        "metni onaylatılmaktadır. Bu metin, şeffaflık ve mevzuata uygun bilgilendirme "
        "ilkeleriyle uyumludur."
    )

    h(doc, "Uygulama mimarisi ve kullanıcı akışı", 2)
    doc.add_paragraph(
        "Uygulama açılışında Firebase başlatılır; “beni hatırla” tercihi kapalıysa "
        "oturum sonlandırılarak cihaz paylaşımlı kullanım senaryolarında risk "
        "azaltılır. Oturum açıkken soğuk açılışta etkinlik, yemekhane ve profil "
        "verileri için Firestore ön yükleme yapılabilir; bu sayede ana ekrana "
        "geçiş süresi kısaltılır. Navigasyon, alt sekmeler üzerinden (örneğin ana "
        "sayfa, menü, haberler, geri bildirim, devamsızlık) modüllere erişimi "
        "düzenler."
    )
    doc.add_paragraph(
        "İstemci tarafında HTTP istekleri `http` / `dio` paketleriyle yapılır; "
        "HTML ayrıştırma için `html` paketi kullanılır. Yerel veritabanı işlemleri "
        "`sqflite` ile yürütülür; veritabanı dosyası uygulamanın özel dizininde "
        "tutulur ve diğer uygulamalarca doğrudan okunamaz (işletim sistemi koruması)."
    )

    h(doc, "Modül Bazında İşlevler", 1)
    modules = [
        (
            "Kimlik doğrulama ve kullanıcı profili",
            "E-posta ve parola ile giriş (Firebase Authentication), kayıtta ad, soyad, "
            "öğrenci numarası ve isteğe bağlı profil fotoğrafı. Profil fotoğrafı "
            "Firebase Storage üzerinde barındırılır; kullanıcıya ait temel alanlar "
            "Cloud Firestore `users` koleksiyonunda tutulur. Oturumun cihazda "
            "kalıcılığı `remember_me` tercihi ile yönetilir.",
        ),
        (
            "Ana sayfa ve ders programı",
            "Yerel SQLite veritabanı (`timeTable.db`) üzerinde ders kayıtları; yaklaşan "
            "ders bilgisi ve günlük görünümler. Haftalık program ve bildirimlerle "
            "entegrasyon hedeflenmiştir.",
        ),
        (
            "SIS üzerinden ders programı aktarımı (WebView)",
            "Kullanıcı, gömülü tarayıcıda doğrudan `sis.agu.edu.tr` oturum sayfasında "
            "giriş yapar. Oturum açıldıktan sonra sayfadaki haftalık program tabloları "
            "HTML olarak okunur, yalnızca ders adı, gün, saat, sınıf ve öğretim üyesi "
            "bilgisi ayrıştırılarak yerel veritabanına yazılır. Bu süreç aşağıda iki "
            "ayrı alt başlıkta açıklanmıştır.",
        ),
        (
            "Devamsızlık takibi",
            "Kullanıcının kendi kaydettiği dersler için günlük yoklama benzeri "
            "işaretleme ve yerel devamsızlık sayacı (`lessons` tablosunda attendance). "
            "Veriler öncelikle cihazda tutulur; ilgili ekranlar SQLite üzerinden çalışır.",
        ),
        (
            "Yemekhane (günlük ve aylık menü)",
            "Menü verileri Firestore `refectory_menus` koleksiyonundan okunur; "
            "`app_cache_meta/sync` ile önbellek parmak izi kontrol edilir ve "
            "SharedPreferences üzerinde yerel önbellek tutulur. Böylece çevrimdışı "
            "senaryolarda son başarılı veri gösterilebilir.",
        ),
        (
            "Etkinlikler, konferanslar, geziler ve diğer duyurular",
            "İlgili içerikler Firestore’daki `events`, `speakers`, `trips` koleksiyonlarından "
            "yüklenir; yemekhane verisiyle birlikte toplu çekim modeli kullanılır.",
        ),
        (
            "Haberler",
            "Harici haber kaynağı WebView ile gösterilir; kullanıcı üniversite web "
            "içeriğini uygulama içinde görüntüler.",
        ),
        (
            "Rehber (AGÜ ve Kayseri)",
            "İletişim ve yönlendirme bilgileri; bağlantılar `url_launcher` ile açılabilir.",
        ),
        (
            "Akademik takvim ve menü araçları",
            "Akademik takvim, bildirim tercihleri, geliştirici bilgisi ve geri bildirim "
            "gibi öğeler menü altında toplanmıştır.",
        ),
        (
            "Üniversite platformlarına hızlı erişim",
            "SIS, Canvas ve e-posta (Zimbra) gibi sistemlere yönlendirme butonları; "
            "kullanıcı tercihine göre harici tarayıcı veya uygulama içi görünüm.",
        ),
        (
            "Geri bildirim",
            "Metin, puan ve isteğe bağlı dosya eki. Dosyalar Firebase Storage’da "
            "`feedback_uploads/` altında; özet bilgi Firestore `feedbacks` koleksiyonunda.",
        ),
        (
            "Kablosuz ağ bilgilendirme",
            "Kampüs kablosuz ağlarına ilişkin kullanıcıyı yönlendiren bilgilendirme "
            "ekranı (üniversitenin duyurduğu erişim bilgilerinin mobilde tek yerden "
            "görülebilmesi amacıyla).",
        ),
        (
            "Bildirimler",
            "Yerel bildirim altyapısı (`flutter_local_notifications`); izinler "
            "işletim sistemi düzeyinde istenir.",
        ),
        (
            "İsteğe bağlı bulut senkronizasyonu (ders / sınıf verisi)",
            "Bazı ekranlarda kullanıcıya özel `users/{id}/classes` alt koleksiyonları "
            "kullanılarak haftalık program veya ders özetleri Firestore üzerinden "
            "sunulabilmektedir. Bu yapı, kullanıcı hesabına bağlıdır ve erişim kuralları "
            "Firebase güvenlik modeli ile sınırlandırılmalıdır (üretim ortamında "
            "kuralların üniversite politikasına göre gözden geçirilmesi önerilir).",
        ),
    ]
    for title, desc in modules:
        h(doc, title, 2)
        doc.add_paragraph(desc)

    h(doc, "Kullanıcı Anketi — Bulgular", 1)
    doc.add_paragraph(
        "Aşağıdaki grafikler, uygulama kullanıcılarına yönelik kısa bir anketin "
        "özetidir. Anket, zaman kısıtı nedeniyle hızlı şekilde yürütülmüş; toplam "
        "31 yanıt elde edilmiştir. Örneklemin sınırlılığı bilinçli olarak raporda "
        "belirtilmiştir: yönetim daha geniş katılımlı bir çalışma talep ederse, "
        "benzeri anketler genişletilerek kurula sunulabilir."
    )
    survey_numbers = """
• Mobil uygulama ihtiyacı: Katılıyorum %96,8; Kararsız %0; Katılmıyorum %3,2 (yaklaşık 1 kişi).
• Uygulamanın resmî platform olması ve kontrolün üniversiteye veya seçilmiş öğrenci topluluğuna devri:
  Her iki “katılıyorum” seçeneği toplam %100; “Kararsız” ve “Katılmıyorum” %0.
  Yönetim modeli tercihi: üniversite + öğrenci kurulu esnekliği %64,5; yalnızca üniversite yönetimi %35,5.
• Memnuniyet: Çok memnun %67,7; Memnun %25,8; birlikte olumlu görüş %93,5 civarı; çok düşük oranda kararsız / memnun değil.
• Öğrenci faydası: Katılıyorum %96,8.
• Kullanım sıklığı: Okuldayken sık sık %71; haftada birkaç kez %19,4; çok nadiren ve neredeyse hiç birlikte küçük bir dilim.
• Güvenlik algısı: Uygulamanın güvenilir olduğu görüşü %83,9; emin değilim %12,9; gizlilik/güvenlik endişesi yaklaşık %3,2 (1 kişi).
• En çok kullanılan modüller (çoklu seçim): Yemekhane %93,5; Ders programı %71; Devamsızlık takibi %51,6; SIS’ten otomatik program %41,9; Akademik takvim %38,7; hızlı erişim %32,3; rehber %9,7.
"""
    for line in survey_numbers.strip().split("\n"):
        doc.add_paragraph(line.strip(), style="List Bullet")

    doc.add_paragraph(
        "Anket özetinin görsel teyidi aşağıdaki şekillerdedir (ekran görüntüleri "
        "anket arayüzünden alınmıştır)."
    )
    for fn, cap in SURVEY_IMAGES:
        add_picture_safe(doc, fn, cap)

    h(doc, "Gizlilik ve Güvenlik", 1)
    doc.add_paragraph(
        "Aşağıdaki alt başlıklar, resmî inceleme ve denetim süreçlerinde sıkça "
        "sorulan sorulara doğrudan yanıt verecek şekilde yapılandırılmıştır. "
        "Amaç; hem yönetim hem de teknik ekiplerin ortak bir dilde buluşmasını "
        "sağlamaktır."
    )

    h(doc, "6698 sayılı KVKK ile uyum perspektifi", 2)
    doc.add_paragraph(
        "Veri sorumlusu ve işleyen rolleri, uygulama üniversite ile resmî olarak "
        "entegre edildiğinde üniversite politikası ve sözleşmelerle netleştirilmelidir. "
        "Mevcut aşamada kullanıcıya açık rıza ve bilgilendirme (kayıt ekranı) "
        "sunulmaktadır. İşlenen veri kategorileri: kimlik (ad, soyad, e-posta, "
        "öğrenci numarası), isteğe bağlı görsel (profil fotoğrafı), kullanım amaçlı "
        "geri bildirim metni ve eki dosyalar, yerel ders ve devamsızlık tutumları. "
        "Açık rıza, belirli bir amaç için alınmış olup; amaç değişirse kullanıcıya "
        "yeniden bilgilendirme yapılması esastır."
    )

    h(doc, "Veri sınıflandırması: cihazda mı, bulutta mı?", 2)
    doc.add_paragraph(
        "Cihazda (yerel): SQLite veritabanında ders programı satırları ve devamsızlık "
        "sayacı; SharedPreferences’ta oturum “beni hatırla”, bildirim/önbellek anahtarları "
        "ve etkinlik-yemekhane önbelleği (JSON string olarak, boyut kontrolüyle). "
        "Bu veriler kullanıcının telefonunda kalır; yedekleme/iCloud davranışı işletim "
        "sistemine bağlıdır."
    )
    doc.add_paragraph(
        "Bulutta (Google Firebase): Kimlik doğrulama için e-posta adresi Google’ın "
        "kimlik sisteminde tutulur; parola asla düz metin olarak geliştiricilere "
        "iletilmez ve uygulama kodunda saklanmaz. Firestore’da kullanıcı profili "
        "alanları, geri bildirim kayıtları, menü ve etkinlik içerikleri; Storage’da "
        "profil ve geri bildirim dosyaları barındırılır."
    )

    h(doc, "Firebase ve Google altyapısı", 2)
    doc.add_paragraph(
        "Firebase, Google’ın yönettiği bulut hizmetleridir. Veri merkezleri ve "
        "uyumluluk süreçleri Google’ın kurumsal altyapısına dayanır. Firebase "
        "Authentication ile parolalar tek yönlü özet (hash) ve modern kimlik "
        "standartlarıyla korunur; geliştirici ekibi veya veritabanı yöneticisi "
        "olarak “kullanıcı parolasını okuma” mümkün değildir. Şifre sıfırlama "
        "akışı da Google’ın güvenli e-posta bağlantısı üzerinden yürür."
    )

    h(doc, "Ağ iletişimi ve üçüncü taraflar", 2)
    doc.add_paragraph(
        "Uygulama; HTTPS üzerinden Firebase API’leri, üniversite web siteleri ve "
        "SIS gibi hizmetlere bağlanır. Harici siteler WebView veya sistem tarayıcısı "
        "ile açıldığında, çerez ve oturum yönetimi ilgili sitenin politikasına tabidir."
    )

    h(doc, "Veri minimizasyonu ve amaç sınırlaması", 2)
    doc.add_paragraph(
        "SIS’ten çekilen içerik yalnızca ders programı tablosunun kullanıcıya "
        "zaten görünen kısımlarıyla sınırlı tutulur; transkript, not, kimlik "
        "numarası gibi ek alanlar için tarama yapılmaz. Firebase tarafında "
        "kullanıcı profili, uygulama işlevleri için gerekli alanlarla sınırlıdır. "
        "Geri bildirim ekindeki dosyalar isteğe bağlıdır ve kullanıcı bilinçli "
        "seçim yapmadan yüklenmez."
    )

    h(doc, "Geliştirici ekibinin erişim sınırları", 2)
    doc.add_paragraph(
        "Kaynak kodu projeye erişimi olan geliştiriciler, Firebase konsolunda "
        "yapılandırma yetkisi varsa meta verileri veya içerik özetlerini görebilir; "
        "ancak Firebase Authentication’da saklanan parolalar düz metin olarak "
        "hiçbir arayüzde listelenmez. Bu nedenle “şifreyi ekipte saklama” riski "
        "bulut kimlik hizmetinde ortadan kalkar. Yine de üretim ortamında çok "
        "faktörlü kimlik, IP kısıtlaması ve konsol erişim günlükleri gibi kurumsal "
        "önlemler üniversite IT politikasıyla birlikte değerlendirilmelidir."
    )

    h(doc, "Önerilen kurumsal kontroller (Firestore güvenlik kuralları)", 2)
    doc.add_paragraph(
        "Bulut veritabanında kullanıcıya özel koleksiyonların yalnızca ilgili "
        "kimlik doğrulaması yapılmış kullanıcı tarafından okunup yazılabilmesi "
        "için Firebase Security Rules ile politikalar tanımlanmalıdır. Bu rapor "
        "kural dosyası içermez; ancak resmî entegrasyonda üniversite bilgi işlem "
        "biriminin kuralları gözden geçirmesi ve onaylaması önerilir. Ayrıca "
        "düzenli yedekleme ve anahtar rotasyonu Google tarafında standart süreçlerle "
        "yürütülür."
    )

    h(doc, "SIS’ten Ders Programı Alma — Gizlilik ve Güvenlik", 1)

    h(doc, "Herkes için: ne yapılıyor, neden meşru?", 2)
    doc.add_paragraph(
        "Kullanıcı, uygulama içindeki gömülü tarayıcıda tıpkı cep telefonundaki "
        "Chrome veya Safari’de olduğu gibi kendi SIS kullanıcı adı ve parolasıyla "
        "giriş yapar. Girişten sonra ekranda görünen haftalık ders tablosu, kullanıcının "
        "zaten erişim hakkı bulunduğu bilgilerin aynısıdır. Uygulama, bu tabloyu "
        "“kopyala-yapıştır” benzeri şekilde okuyup yalnızca ders adı, saat, sınıf ve "
        "öğretim üyesi gibi alanları ayıklayarak telefonda yerel listeye yazar. "
        "Başka öğrencilerin verisine erişim, sunucuya izinsiz bağlantı veya veri "
        "“sızdırma” söz konusu değildir; işlem, kullanıcının kendi oturumuyla "
        "gördüğü sayfanın düzenlenmesidir."
    )

    h(doc, "Teknik açıdan: WebView, JavaScript kancaları ve ayrıştırma", 2)
    doc.add_paragraph(
        "Teknik olarak WebView denetleyicisi, sayfa yüklendikten sonra iframe "
        "içeriğinin yüklenmesini izleyen küçük betikler (hook) çalıştırır; haftalık "
        "program HTML’i `window.__SIS_IFR_READY_HTML` gibi geçici bir değişkende "
        "toplanır. `html` paketi ile DOM ayrıştırılır; `grd0`–`grd4` kimlikli "
        "tablolardan veya yedek olarak genel tablolardan satırlar okunur. Aynı "
        "içerik tekrar geldiğinde SHA-256 ile özet alınarak gereksiz tekrar yazımı "
        "engellenir. Bu süreç yalnızca istemci tarafında çalışır; SIS’e özel bir "
        "API anahtarı veya gizli bir arka kapı eklenmez. Dio çerez yöneticisi "
        "projede tanımlıdır; oturum çerezleri tarayıcı oturumuyla uyumludur."
    )

    h(doc, "Yayın Durumu: Google Play ve Apple App Store", 1)
    doc.add_paragraph(
        "Uygulamanın Android sürümü Google Play’de yayımlanmıştır. Geçmişte, "
        "süreç ve kurumsal iletişim hakkında yeterince bilgi sahibi olunmadan "
        "yayımlanmış olması nedeniyle üniversite yönetiminin haberdar edilmemesi "
        "kusur teşkil etmiştir. Ekip, bu durumdan dolayı mahcubiyet duymakta; "
        "bürokratik ve kurumsal onay mekanizmalarını yeterince öngörememiş "
        "olmanın verdiği aceleciliği samimiyetle kabul etmektedir."
    )
    doc.add_paragraph(
        "iOS sürümü teknik olarak App Store’a çıkmaya hazırdır; ancak ekip, "
        "üniversite ile mutabakat sağlanmadan yeni bir kamusal yayımlama adımı "
        "atmamayı ilke edinmiştir. Böylece kurumsal itibar ve tek kanaldan "
        "yürütülecek iletişim ön planda tutulmaktadır."
    )

    h(doc, "Gelecekte Genişletilebilecek Modüller", 1)
    doc.add_paragraph(
        "Üniversite onayı ve veri entegrasyonu sağlandığında aşağıdaki gibi modüller "
        "geliştirilebilir (örnekler):"
    )
    futures = [
        "Spor salonu ve spor tesisleri için randevu / rezervasyon sistemi.",
        "Kütüphane oturumu veya kaynak rezervasyonu.",
        "Akademisyenlerin ilgili dersi alan öğrencilere yönelik duyuru ve mesajlaşma "
        "kanalı (rol tabanlı yetkilendirme ve KVKK uyumu ile).",
    ]
    for item in futures:
        doc.add_paragraph(item, style="List Bullet")

    h(doc, "Teslim, Kontrol ve Resmîleşme Niyeti", 1)
    doc.add_paragraph(
        "Geliştirici ekibin temel niyeti; uygulamanın işlevsel olarak tamamlanması "
        "ve üniversite yönetimine, kontrol ve yönetimin tamamen okula ait olacağı "
        "şekilde devredilmesidir. Üniversitenin uygun görmesi halinde uygulamanın "
        "Abdullah Gül Üniversitesi’nin resmî dijital platformlarından biri haline "
        "getirilmesi; marka, içerik, yasal metinler ve operasyonun kurumsal süreçlere "
        "bağlanması hedeflenmektedir."
    )

    h(doc, "Riskler, sınırlılıklar ve iyileştirme alanları", 1)
    doc.add_paragraph(
        "Her mobil uygulama gibi AGÜ Mobile da belirli sınırlılıklara sahiptir. "
        "Bunlar raporda gizlenmemeli; kurumsal güven için açıkça listelenmelidir: "
        "(1) Anket örneklemi küçüktür; istatistiksel genelleme için yeterli "
        "sayıda katılımcı hedeflenmelidir. (2) WebView tabanlı SIS entegrasyonu, "
        "üniversite arayüzü değişirse ayrıştırma kodunun güncellenmesini gerektirebilir. "
        "(3) Firebase güvenlik kuralları üniversite onayıyla sıkılaştırılmalıdır. "
        "(4) Kullanıcıların geri bildirim ekinde kişisel veri paylaşmaması için "
        "uyarı metinleri güçlendirilebilir."
    )

    h(doc, "Sonuç", 1)
    doc.add_paragraph(
        "AGÜ Mobile; öğrenci deneyimini iyileştirmeyi hedefleyen, şeffaf bir güvenlik "
        "modeli ve modüler yapı sunan bir istemcidir. Küçük örneklemli anket, "
        "yönelimin güçlü olduğunu göstermektedir; daha geniş örneklemli çalışmalar "
        "isteğe bağlı olarak planlanabilir. Üniversite yönetiminin görüş ve onayı "
        "doğrultusunda yol haritası birlikte netleştirilebilir."
    )

    h(doc, "Ek: Kaynaklar ve Bağlantılar", 1)
    doc.add_paragraph(
        "SIS oturum adresi (kullanıcı arayüzü): https://sis.agu.edu.tr/oibs/std/login.aspx"
    )
    doc.add_paragraph(
        "Firebase güvenlik ve kimlik dokümantasyonu: https://firebase.google.com/support/privacy"
    )

    # Son sayfa: üretim notu
    doc.add_page_break()
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "Bu belge, python-docx ile otomatik üretilmiştir.\n"
        "İçindekiler alanını Word’de güncellemeyi unutmayınız."
    )
    fr.font.size = Pt(8)
    fr.italic = True

    return doc


def main() -> None:
    doc = build()
    doc.save(OUT_PATH)
    print(f"Yazıldı: {OUT_PATH}")


if __name__ == "__main__":
    main()
